import cv2  
from mistralai import Mistral
import pytesseract
import numpy as np
import os
from docx import Document
#from dotenv import load_dotenv
#from langchain.prompts import PromptTemplate #incase of further optimizing this agent we can use langchain or if the guideline doc is very big also


pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"



#opening the fastener guidelines file and reading file
doc_path = r"F:\intern Synera AI\intern-ai\Fastener_Types_Manual.docx"
doc = Document(doc_path)
guidelines = "\n".join([para.text for para in doc.paragraphs])

#opening machine guidelines file and reading file
machine_doc_path = r"F:\intern Synera AI\intern-ai\Manufacturing Expert Manual.docx"
machine_doc = Document(machine_doc_path)
machine_guidelines = "\n".join([para.text for para in machine_doc.paragraphs])


#Open AI requesting and API
# load_dotenv()
# mistral_api_key = os.getenv("MISTRAL_API_KEY")
# client = Mistral(api_key = mistral_api_key )

client = Mistral(api_key="Cc8p5DVZUnwOnRk5C6YekMhcQTTHXGIk")


#function for image analysis
def image_analysis(image_path):
    image = cv2.imread(image_path)

    # Check if the image loaded correctly
    if image is None:
        print("Error: Could not load image. Please check the file path!")
        return None
    
    else:
        #converting to gray scale
        gray = cv2.cvtColor(image,cv2.COLOR_BGR2GRAY)

        #extract text from gray image
        text = pytesseract.image_to_string(gray)

        #edge detection
        edges = cv2.Canny(gray, 50, 150)
        edge_count = np.count_nonzero(edges)

        # Contour detection
        contours, _ = cv2.findContours(edges, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
        num_contours = len(contours)

        # Bounding Box & Aspect Ratio
        if num_contours > 0:
            x, y, w, h = cv2.boundingRect(contours[0])
            aspect_ratio = round(float(w) / h, 2) if h != 0 else 0
        else:
            aspect_ratio = None

        # Convexity & Solidity (How "solid" the object is)
        solidity = None
        if num_contours > 0:
            area = cv2.contourArea(contours[0])
            hull = cv2.convexHull(contours[0])
            hull_area = cv2.contourArea(hull)
            solidity = round(float(area) / hull_area, 2) if hull_area != 0 else None

        # Color Histogram (for material detection)
        avg_color = cv2.mean(image)[:3]  # Average BGR color


        # Display the image in a new window
        cv2.imshow("Edges", edges)
        print("\nNumber of edges:", edge_count)

        # Wait for a key press/2 sec and close the window
        cv2.waitKey(2000)
        cv2.destroyAllWindows()

        image_description = (
            f"Detected text: {text.strip()},"
            f"Edge count: {edge_count},"
            f"Contours detected: {num_contours}, "
            f"Aspect ratio: {aspect_ratio}, "
            f"Solidity: {solidity}, "
            f"Average color (BGR): {avg_color}."
        )

        return image_description
    

# Function to get fastener recommendations
def get_fastener_recommendation(image_description, guidelines,machine_guidelines):
    prompt = f"""You are an AI assistant that provides fastener and machine recommendations.
    Fastener recommendation guidelines:
    {guidelines}
    machine recommendation guidelines:
    {machine_guidelines}
    Image analysis for fastner recomendation:
    {image_description}

    Based on the guidelines and image analysis, suggest:
    1. The most suitable fastener type.
    2. The best manufacturing machine for this geometry.
    Explain your reasoning clearly.

    """

    
    print("\nSending request to Mistral AI...\n")  # Debugging print
    response = client.chat.complete(
        model="mistral-small",  
        messages=[{"role": "user", "content": prompt}]
    )
    print("\nResponse received!\n")  # Debugging print
    return response.choices[0].message.content

    
# Load the image folder
image_folder = r"F:\intern Synera AI\intern-ai\Sample Images"

#creating a list of .png files
image_files = [f for f in os.listdir(image_folder) if f.endswith(".png")]

#printing option
print("\nAvailable Images for analysis:")
for idx, file in enumerate(image_files):
    print(f"{idx}. {file}")


#asking user for the index of image he wishes to analyse
choose_again = True
while choose_again:
    try:
        choice = int(input("\nEnter the index number of the image you want to analyze from the available list: ")) 
        if 0 <= choice < len(image_files):
            selected_image = os.path.join(image_folder, image_files[choice])
            choose_again = False
        else:
            print("You have input an Invalid index")
    except ValueError:
        print("Invalid input! Please enter an index number.")


#calling image analysis function
if image_analysis != None:
    image_description = image_analysis(selected_image)

    

answer = get_fastener_recommendation(image_description, guidelines,machine_guidelines)

print(f"\nSuitbale Recommendations for {selected_image}:\n {answer}")
print("\n"*10)